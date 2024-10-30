from docx import Document

# 创建文档
doc = Document()

# 添加标题
doc.add_heading('飞行报告', 0)

# 添加基本信息
doc.add_paragraph(f'航班号: {flight_data["航班号"]}')
doc.add_paragraph(f'起飞时间: {flight_data["起飞时间"]}')
doc.add_paragraph(f'降落时间: {flight_data["降落时间"]}')

# 添加飞行过程表格
doc.add_paragraph('飞行过程:')
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '时间'
hdr_cells[1].text = '高度'
hdr_cells[2].text = '备注'

for process in flight_data['飞行过程']:
    row_cells = table.add_row().cells
    row_cells[0].text = process['时间']
    row_cells[1].text = process['高度']
    row_cells[2].text = process['备注']

# 保存文档
doc.save('flight_report.docx')