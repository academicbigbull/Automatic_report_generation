from docx import Document
from docx.shared import Inches

import win32com#win32com（主要用作doc转docx格式转换用）
from win32com.client import Dispatch, constants

from mailmerge import MailMerge#docx-mailmerge（用作按照模板生成大量同类型文档）
import matplotlib.pyplot as plt


#创建一个word文档
# 导包
from docx import Document

# 实例化一个Document对象，相当于打开word软件，新建一个空白文件
doc = Document()
# word文件尾部增加一个段落，并写入内容
paragraph = doc.add_paragraph('Hello! I am created by python-docx.')
# 原有段落前面插入一个段落，并写入内容
prior_paragraph = paragraph.insert_paragraph_before('Good day!')
# 保存word文件到当前文件夹
doc.save(r"E:\python\自动生成word图文报告\test.docx")


#------------------------------------------------------------------------------
# 增加标题：add_heading(self, text="", level=1):
doc.add_heading('标题',1)

# 增加段落正文，add_paragraph(self, text='', style=None)：返回一个 Paragraph 段落对象
par = doc.add_paragraph('第一个段落：')
print(par)  # <docx.text.paragraph.Paragraph object at 0x000000000A889F08>

# 在段落中添加文字块，add_run(self, text=None, style=None):返回一个 run 对象
run_ = par.add_run('段落文字块')
print(run_)     # <docx.text.run.Run object at 0x000000000B2D31C8>

# 增加图片，add_picture(self, image_path_or_stream, width=None, height=None):返回一个 InlineShape 对象
pic_par = doc.add_picture(r"C:\Users\cm\Pictures\Saved Pictures\A68FDC385F6CAF1CCB92F11C7A2B9828.png")
print(pic_par)  # <docx.shape.InlineShape object at 0x000000000B2F11C8>

p2 = doc.add_paragraph('第二个段落：')
p2.add_run('段落文字块')

# 增加表格add_table(self, rows, cols, style=None):返回一个表格对象
table = doc.add_table(2,3)
print(table)  # <docx.table.Table object at 0x000000000B302688>

# 保存文件
doc.save('test.docx')

#---------------------------------------------------------------
# 增加分页符
doc.add_page_break()

# 增加标题 API 分析， 只能设置 0-9 级标题
for i in range(0,10):
    doc.add_heading(f'标题{i}', i)
doc.save(r'增加标题.docx')
#-------------------------------------------------------------------------

#内容修改，正则表达式替换原有表达式的英文引号
from docx.oxml.ns import qn # 设置中文字体需导入 qn 模块
import re

from docx import Document
from docx.oxml.ns import qn  # 设置中文字体需导入 qn 模块
import re

# 加载Word文档
doc = Document(r"正则替换test.docx")

# 定义正则表达式模式，匹配所有英文双引号中的文本
restr = '"(?:[^"])*"'

# 遍历文档中的所有段落
for p in doc.paragraphs:
    # 查找所有符合模式的文本（即被英文双引号包围的文本）
    list_results = re.findall(restr, p.text)

    # 遍历查找到的所有文本
    for result in list_results:
        # 替换段落中的英文双引号为中文书名号，并保留其中的文字
        p.text = p.text.replace(result, '“' + result[1:-1] + '”')

    # 修改格式的操作必须放在替换操作之后，否则替换操作的效果会被覆盖
    # 遍历段落中的所有run（即连续的相同格式的文字）
    for run in p.runs:
        # 设置run中的西文字体为 `Times New Roman`
        run.font.name = 'Times New Roman'

        # 使用 `qn` 函数设置东亚语言（如中文）的字体为 `宋体`
        run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 保存修改后的文档
doc.save(r'test_modify.docx')


#--------------------------------------------------------------------



