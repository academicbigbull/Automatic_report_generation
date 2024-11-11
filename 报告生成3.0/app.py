from flask import Flask, request, send_file, render_template, abort
from docx import Document
from docx.shared import Pt, Inches,RGBColor
from docx.oxml.ns import qn
from xml.etree import ElementTree as ET
from datetime import datetime
import pandas as pd
import matplotlib.pyplot as plt
import json
import io
import os
import textwrap
from docx.shared import RGBColor  # 导入颜色设置所需的库

app = Flask(__name__)


@app.route('/')
def index():
    return render_template('实例页面.html')


@app.route('/generate', methods=['POST'])
def generate_report():
    try:
        # 获取用户的样式和内容设置
        title_font = request.form.get('titleFont', '宋体')
        title_size = int(request.form.get('titleSize', 12))
        reporter = request.form.get('reporter', '未知报告人')
        generate_summary = request.form.get('generateSummary') == 'true'
        generate_table = request.form.get('generateTable') == 'true'
        generate_image = request.form.get('generateImage') == 'true'

        # 处理上传文件
        data_file = request.files.get('dataFile')
        if data_file:
            file_content = data_file.read()  # 确保文件内容已读取
            if data_file.filename.endswith('.json'):
                data = json.loads(file_content)
            elif data_file.filename.endswith('.xml'):
                tree = ET.ElementTree(ET.fromstring(file_content.decode()))
                root = tree.getroot()
                data = parse_xml_to_list(root)  # 使用新的解析函数
            else:
                return "Unsupported file format", 400

        # 预处理数据（将 JSON 或 XML 文件的数据转成 DataFrame）
        if isinstance(data, list):
            df = pd.DataFrame(data)
        elif isinstance(data, dict):
            df = pd.json_normalize(data)
        else:
            abort(400, description="Invalid data format")

        # 生成 Word 文档
        doc = Document()

        # 应用用户的样式设置，生成标题
        title = doc.add_heading('报告', level=1)
        title_run = title.runs[0]
        title_run.font.size = Pt(title_size)  # 设置字体大小
        title_run.font.name = title_font  # 设置字体名称
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)  # 设置中文字体
        # 应用颜色设置
        title_color = request.form.get('titleColor', '#000000')  # 获取颜色，默认为黑色
        rgb_color = RGBColor(int(title_color[1:3], 16), int(title_color[3:5], 16), int(title_color[5:], 16))
        title_run.font.color.rgb = rgb_color  # 设置标题颜色
        title.alignment = 1  # 居中对齐

        # 添加作者并设置右对齐
        author_paragraph = doc.add_paragraph(f'报告人: {reporter}')
        author_paragraph.alignment = 2  # 2表示右对齐，0表示左对齐，1表示居中对齐

        # 生成摘要
        if generate_summary:
            # 添加标题并设置为黑色
            summary_title = doc.add_heading('摘要', level=2)
            summary_title.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色，RGB值为(0, 0, 0)
            summary_text = generate_summary_text(df)
            doc.add_paragraph(summary_text)

        # 生成表格
        if generate_table:
            # 添加标题并设置为黑色
            table_title = doc.add_heading('数据表格', level=2)
            table_title.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色，RGB值为(0, 0, 0)
            table = doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells = table.rows[0].cells
            for i, column in enumerate(df.columns):
                hdr_cells[i].text = column

            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

        # 生成特征图表
        if generate_image:
            # 添加标题并设置为黑色
            image_title = doc.add_heading('图表', level=2)
            image_title.runs[0].font.color.rgb = RGBColor(0, 0, 0)  # 设置颜色为黑色，RGB值为(0, 0, 0)
            img_stream = generate_feature_chart(df)
            if img_stream:
                doc.add_picture(img_stream, width=Inches(5))
                doc.paragraphs[-1].alignment = 1  # 图片居中

        # 保存文档到内存
        filename = f'report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        doc.save(filename)
        with open(filename, 'rb') as f:
            file_data = f.read()
        os.remove(filename)

        # 返回生成的文件
        return send_file(io.BytesIO(file_data), as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        print(f"An error occurred: {e}")
        abort(500,
              description="An error occurred while generating the report. Please check the input data and try again.")


def parse_xml_to_list(root):
    """解析 XML 文件，将每个 Employee 数据提取为字典并返回列表。"""
    data = []
    for element in root.findall("Employee"):
        employee_data = {}
        for child in element:
            employee_data[child.tag] = child.text
        data.append(employee_data)
    return data


def generate_summary_text(df):
    """生成简化版的摘要文本，不依赖任何外部分词器。"""
    text = df.to_string(index=False)  # 将 DataFrame 转换为字符串
    summary_text = "\n".join(textwrap.wrap(text, width=80)[:5])  # 取前5行作为摘要
    return summary_text if summary_text else "数据量太小，无法生成摘要。"


def generate_feature_chart(df):
    """生成图表并保存到内存。"""
    img_stream = io.BytesIO()
    try:
        plt.rcParams['font.sans-serif'] = ['SimHei']  # 使用黑体
        plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题

        # 尝试将所有列转换为数值，非数值列将变为 NaN
        df = df.apply(pd.to_numeric, errors='coerce')

        # 仅保留数值列并去掉全 NaN 列
        numeric_df = df.select_dtypes(include=['float64', 'int64']).dropna(axis=1, how='all')
        if numeric_df.empty:
            raise ValueError("No numeric data to plot")

        numeric_df.plot(kind='bar', figsize=(10, 6))  # 使用数值列绘图
        plt.title("数据特征图表")
        plt.xlabel("Index")
        plt.ylabel("Values")
        plt.tight_layout()
        plt.savefig(img_stream, format='png')
        img_stream.seek(0)
    except Exception as e:
        print(f"Error generating chart: {e}")
        img_stream = None
    finally:
        plt.close()
    return img_stream


if __name__ == '__main__':
    app.run(debug=True)
