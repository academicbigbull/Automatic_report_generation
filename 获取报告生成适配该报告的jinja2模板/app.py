from docx import Document
from jinja2 import Template


def extract_docx_features(file_path):
    # 读取 docx 文件
    doc = Document("E://python//获取报告生成适配该报告的jinja2模板//rest_1.docx")

    # 初始化特征字典
    features = {
        "title": None,
        "title_color": None,  # 标题颜色
        "title_size": None,  # 标题字体大小
        "report_people": None,  # 汇报人
        "report_people_color": None,  # 汇报人字体颜色
        "report_people_size": None,  # 汇报人字体大小
        "report_datatime": None,  # 实验日期
        "report_datatime_color": None,  # 实验日期字体颜色
        "report_datatime_size": None,  # 实验日期字体大小
        "report_place": None,  # 实验地点
        "report_place_color": None,  # 实验地点字体颜色
        "report_place_size": None,  # 实验地点字体大小
        # 表格和图片不用设置格式
        "data_chart": [],
        "data_picture": None
    }

    # 提取标题信息（假设第一段为标题）
    if doc.paragraphs:
        title_paragraph = doc.paragraphs[0]
        features["title"] = title_paragraph.text
        features["title_color"] = title_paragraph.style.font.color.rgb if title_paragraph.style.font.color else None
        features["title_size"] = title_paragraph.style.font.size.pt if title_paragraph.style.font.size else None

    # 提取汇报人、实验日期、实验地点等信息
    for para in doc.paragraphs:
        if '汇报人' in para.text:
            features["report_people"] = para.text.split('：')[1].strip()  # 假设格式是 "汇报人：XXX"
            features["report_people_color"] = para.style.font.color.rgb if para.style.font.color else None
            features["report_people_size"] = para.style.font.size.pt if para.style.font.size else None
        elif '实验日期' in para.text:
            features["report_datatime"] = para.text.split('：')[1].strip()  # 假设格式是 "实验日期：YYYY-MM-DD"
            features["report_datatime_color"] = para.style.font.color.rgb if para.style.font.color else None
            features["report_datatime_size"] = para.style.font.size.pt if para.style.font.size else None
        elif '实验地点' in para.text:
            features["report_place"] = para.text.split('：')[1].strip()  # 假设格式是 "实验地点：XXX"
            features["report_place_color"] = para.style.font.color.rgb if para.style.font.color else None
            features["report_place_size"] = para.style.font.size.pt if para.style.font.size else None

    # 提取表格内容
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        features["data_chart"].append(table_data)

    # 提取图片信息（获取图片的路径或URL）
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            features["data_picture"] = rel.target_ref  # 提取图片路径或URL

    return features


def generate_jinja2_template(features):
    # 生成动态的 Jinja2 模板
    template = """
    <!DOCTYPE html>
    <html lang="zh">
    <head>
        <meta charset="UTF-8">
        <title>{{ title }}</title>
        <style>
            body {
                font-family: "Arial", sans-serif;
            }
            h1 {
                font-size: {{ title_size }}pt;
                color: {{ title_color }};
            }
            .report-people {
                font-size: {{ report_people_size }}pt;
                color: {{ report_people_color }};
            }
            .report-datetime {
                font-size: {{ report_datatime_size }}pt;
                color: {{ report_datatime_color }};
            }
            .report-place {
                font-size: {{ report_place_size }}pt;
                color: {{ report_place_color }};
            }
            table {
                border-collapse: collapse;
                width: 100%;
            }
            th, td {
                border: 1px solid black;
                padding: 8px;
                text-align: center;
            }
        </style>
    </head>
    <body>
        {% if title %}
        <h1>{{ title }}</h1>
        {% endif %}

        {% if report_people %}
        <p class="report-people">汇报人: {{ report_people }}</p>
        {% endif %}

        {% if report_datatime %}
        <p class="report-datetime">实验日期: {{ report_datatime }}</p>
        {% endif %}

        {% if report_place %}
        <p class="report-place">实验地点: {{ report_place }}</p>
        {% endif %}

        {% if data_chart %}
        <h2>实验数据</h2>
        {% for table in data_chart %}
            <table>
                <thead>
                    <tr>
                        {% for column in table[0] %}
                            <th>{{ column }}</th>
                        {% endfor %}
                    </tr>
                </thead>
                <tbody>
                    {% for row in table[1:] %}
                        <tr>
                            {% for cell in row %}
                                <td>{{ cell }}</td>
                            {% endfor %}
                        </tr>
                    {% endfor %}
                </tbody>
            </table>
        {% endfor %}
        {% endif %}

        {% if data_picture %}
        <h2>报告图片</h2>
        <img src="{{ data_picture }}" alt="报告图片" />
        {% endif %}

    </body>
    </html>
    """

    return template


# 使用示例
file_path = "E://python//获取报告生成适配该报告的jinja2模板//rest_1.docx"
report_features = extract_docx_features(file_path)

# 生成 Jinja2 模板
template = generate_jinja2_template(report_features)

# 使用 Jinja2 渲染模板
jinja_template = Template(template)
rendered_html = jinja_template.render(report_features)

# 打印渲染后的 HTML
print(rendered_html)
