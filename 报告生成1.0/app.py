from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
import pandas as pd
import io
import os

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index1.html')

@app.route('/generate', methods=['POST'])
def generate_report():
    try:
        # 获取表单数据
        name = request.form['name']
        date = request.form['date']
        data = request.form['data']
        image_file = request.files['image']

        # 将字符串数据转换为 JSON 格式
        import json
        data = json.loads(data)

        # 处理数据
        df = pd.DataFrame(data)

        # 创建一个新的 Word 文档
        doc = Document()

        # 添加标题
        title = doc.add_heading('报告', level=1)
        title.alignment = 1  # 居中对齐

        # 添加作者和日期
        doc.add_paragraph(f'作者: {name}')
        doc.add_paragraph(f'日期: {date}')

        # 添加表格
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(df.columns):
            hdr_cells[i].text = column

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        doc.add_paragraph()#表格中间与图片空行

        # 如果有上传的图片，则插入图片并设置居中
        if image_file:
            image_format = image_file.filename.split('.')[-1].lower()
            if image_format in ['jpg', 'jpeg', 'png']:
                image_stream = io.BytesIO(image_file.read())
                picture = doc.add_picture(image_stream, width=Inches(4))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = 1
            else:
                return "Unsupported image format. Please upload a JPG or PNG image.", 400

        # 保存文档到内存中
        filename = f'report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        doc.save(filename)
        with open(filename, 'rb') as f:
            file_data = f.read()
        os.remove(filename)

        # 返回生成的文件，设置适当的响应头和文件名
        return send_file(io.BytesIO(file_data), as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        # 打印错误信息并返回一个友好的错误消息给用户
        print(f"An error occurred: {e}")
        return "An error occurred while generating the report.", 500

if __name__ == '__main__':
    app.run(debug=True)
