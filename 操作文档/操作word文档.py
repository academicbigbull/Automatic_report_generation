from docx import Document
from docx.shared import Cm,Pt,RGBColor#Cm模块用于设定图片的尺寸大小,pt设置字体大小或其他需要以磅（points）为单位的属性,RGBColor用于设置颜色，特别是字体颜色或背景色。RGBColor 类接受三个参数，分别表示红、绿、蓝三种颜色的强度值（0-255）。
from openpyxl import Workbook
from docx.oxml.ns import qn

#读取文档的文字
doc = Document(r'E:/python/操作文档/test_word.docx')
print(doc.paragraphs)
# for paragraph in doc.paragraphs:
#     print(paragraph.text)

#读取文字块
paragraph=doc.paragraphs[0]
runs=paragraph.runs
print(runs)
for run in paragraph.runs:
    print(run.text)
#------------------------------
paragraph = doc.paragraphs[1]
runs = paragraph.runs
print(runs)
for run in paragraph.runs:
    print(run.text)

#添加文字
# paragraph1=doc.add_paragraph('这是添加的第一段')
# paragraph1=doc.add_paragraph('这是添加的第二段')
# doc.save(r'E:/python/操作文档/test_word.docx')

#增添文字块
# paragraph3 = doc.add_paragraph()
# paragraph3.add_run("我被加粗了文字块儿").bold = True
# paragraph3.add_run("，我是普通文字块儿，")
# paragraph3.add_run("我是斜体文字块儿").italic = True
# doc.save(r'E:/python/操作文档/test_word.docx')

#添加一个分页
# doc.add_page_break()
# doc.save(r'E:/python/操作文档/test_word.docx')

#添加图片
# doc.add_picture(r"C:\Users\cm\Pictures\Saved Pictures\A68FDC385F6CAF1CCB92F11C7A2B9828.png",width=Cm(5),height=Cm(5))
# doc.save(r'E:/python/操作文档/test_word.docx')

#添加表格
# list1 = [
#     ["姓名", "性别", "家庭地址"],
#     ["唐僧", "男", "湖北省"],
#     ["孙悟空", "男", "北京市"],
#     ["猪八戒", "男", "广东省"],
#     ["沙和尚", "男", "湖南省"]
# ]
# list2 = [
#     ["姓名", "性别", "家庭地址"],
#     ["貂蝉", "女", "河北省"],
#     ["杨贵妃", "女", "贵州省"],
#     ["西施", "女", "山东省"]
# ]
#
# table1 = doc.add_table(rows=5, cols=3)
# for row in range(5):
#     cells = table1.rows[row].cells
#     for col in range(3):
#         cells[col].text = str(list1[row][col])
# doc.add_paragraph("-----------------------------------------------------------")
# table2 = doc.add_table(rows=4, cols=3)
# for row in range(4):
#     cells = table2.rows[row].cells
#     for col in range(3):
#         cells[col].text = str(list2[row][col])
#
# doc.save(r'E:/python/操作文档/test_word.docx')

#将word的表转存为excel
# t0 = doc.tables[0]
#
# workbook = Workbook()
# sheet = workbook.active
#
# for i in range(len(t0.rows)):
#     list1 = []
#     for j in range(len(t0.columns)):
#         list1.append(t0.cell(i, j).text)
#     sheet.append(list1)
# workbook.save(filename=r"E:/python/操作文档/来自word中的表.xlsx")

