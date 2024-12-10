from flask import Flask, render_template, request, send_file, abort
from xml.etree.ElementTree import ElementTree, fromstring
import json
import pandas as pd
from collections import Counter
from docx import Document
from docx.shared import Pt, RGBColor

app = Flask(__name__)

# 解析 XML 数据
def parse_xml(file_content):
    root = ElementTree(fromstring(file_content)).getroot()
    data = []
    for item in root.findall("item"):
        row = {child.tag: child.text for child in item}
        data.append(row)
    return data

# 解析 JSON 数据
def parse_json(file_content):
    return json.loads(file_content)

# 计算类别分布
def calculate_distribution(data, columns):
    distributions = {}
    for col in columns:
        column_data = [row[col] for row in data]
        distributions[col] = dict(Counter(column_data))
    return distributions

# 生成报告文件
def create_word_report(title_font, title_size, title_color, reporter, generate_summary, summary_text,
                       generate_table, table_data, distributions, total_records, total_fields):
    doc = Document()
    title = doc.add_heading('报告', level=1)
    title_run = title.runs[0]
    title_run.font.size = Pt(title_size)
    title_run.font.name = title_font
    title_run.font.color.rgb = RGBColor(*[int(title_color[i:i+2], 16) for i in (1, 3, 5)])
    title.alignment = 1  # 居中

    doc.add_paragraph(f"报告人: {reporter}")

    if generate_summary:
        doc.add_heading('摘要', level=2)
        for field, distribution in distributions.items():
            doc.add_paragraph(f"【{field}】的类别分布：")
            for key, value in distribution.items():
                doc.add_paragraph(f"  {key}: {value} 个记录")
        doc.add_paragraph(f"本次数据集包含 {total_records} 条记录，涵盖了 {total_fields} 个字段。")

    if generate_table:
        doc.add_heading('数据表格', level=2)
        table = doc.add_table(rows=1, cols=len(table_data.columns))
        for i, column in enumerate(table_data.columns):
            table.cell(0, i).text = column
        for row in table_data.itertuples(index=False):
            table_row = table.add_row()
            for i, value in enumerate(row):
                table_row.cells[i].text = str(value)

    return doc

@app.route('/')
def index():
    return render_template('实例页面.html')

@app.route('/generate', methods=['POST'])
def generate_report():
    try:
        # 获取用户输入
        title_font = request.form['titleFont']
        title_size = int(request.form['titleSize'])
        title_color = request.form.get('titleColor', '#000000')
        reporter = request.form['reporter']
        generate_summary = request.form.get('generateSummary') == 'true'
        generate_table = request.form.get('generateTable') == 'true'

        # 上传文件处理
        data_file = request.files['dataFile']
        file_content = data_file.read().decode()
        if data_file.filename.endswith('.json'):
            data = parse_json(file_content)
        elif data_file.filename.endswith('.xml'):
            data = parse_xml(file_content)
        else:
            abort(400, "Unsupported file format")

        # 转换数据为 DataFrame
        table_data = pd.DataFrame(data)

        # 计算类别分布
        distributions = calculate_distribution(data, table_data.columns)

        # 渲染模板生成 HTML
        table_columns = table_data.columns.tolist()
        table_data_values = table_data.values.tolist()
        total_records = len(data)
        total_fields = len(table_data.columns)
        report_html = render_template('report_template.jinja2',
                                      title_font=title_font, title_size=title_size, title_color=title_color,
                                      title_text="报告", reporter=reporter, generate_summary=generate_summary,
                                      summary_text="数据集统计", generate_table=generate_table,
                                      table_columns=table_columns, table_data=table_data_values,
                                      id_distribution=distributions.get('ID', {}),
                                      name_distribution=distributions.get('Name', {}),
                                      age_distribution=distributions.get('Age', {}),
                                      department_distribution=distributions.get('Department', {}),
                                      salary_distribution=distributions.get('Salary', {}),
                                      total_records=total_records, total_fields=total_fields)

        # 将生成的 HTML 转为 PDF 或 Word 文件
        doc = create_word_report(title_font, title_size, title_color, reporter, generate_summary, "数据集统计",
                                 generate_table, table_data, distributions, total_records, total_fields)
        filename = f"{reporter}_report.docx"
        doc.save(filename)

        return send_file(filename, as_attachment=True)

    except Exception as e:
        return f"生成报告出错: {e}", 500

if __name__ == '__main__':
    app.run(debug=True)
